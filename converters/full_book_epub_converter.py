#!/usr/bin/env python3
"""
Full Book EPUB Converter for "Caught in the Act"
Generates complete EPUB with all chapters and enhanced references
"""

import json
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, List
from docx import Document

try:
    from ebooklib import epub
    EBOOKLIB_AVAILABLE = True
except ImportError:
    EBOOKLIB_AVAILABLE = False
    print("⚠️  Warning: ebooklib not installed")


class FullBookEPUBConverter:
    """Convert complete manuscript to EPUB with enhanced references"""

    def __init__(self, docx_path: str, references_json: str = 'data/references_master.json'):
        if not EBOOKLIB_AVAILABLE:
            raise ImportError("ebooklib required. Install: pip install ebooklib")

        print(f"📚 Loading complete manuscript from {docx_path}...")
        self.doc = Document(docx_path)

        print(f"📖 Loading enhanced references from {references_json}...")
        with open(references_json, 'r', encoding='utf-8') as f:
            ref_data = json.load(f)

        self.references = {ref['number']: ref for ref in ref_data['references']}
        self.statistics = ref_data['statistics']

        print(f"✓ Loaded {len(self.doc.paragraphs)} paragraphs from manuscript")
        print(f"✓ Loaded {len(self.references)} enhanced references")

    def convert(self, output_path: str = 'outputs/complete_book.epub'):
        """Convert complete book to EPUB"""
        print(f"\n📚 Converting complete book to EPUB...")

        # Extract structure
        structure = self._extract_structure()

        # Create EPUB
        book = epub.EpubBook()

        # Metadata
        book.set_identifier('caught-in-the-act-complete-001')
        book.set_title('Caught in the Act: Understanding Political Performance in the Digital Age')
        book.set_language('en')
        book.add_author('Joe Ferguson')
        book.add_metadata('DC', 'description',
                          'Complete book with enhanced references and verified URLs')

        # Create CSS
        css = self._create_stylesheet()
        book.add_item(css)

        # Build book chapters
        chapters = []
        toc_entries = []

        # Title page
        title_page = self._create_title_page()
        book.add_item(title_page)
        chapters.append(title_page)

        # Foreword
        if structure['foreword']:
            foreword = self._create_section('Foreword', 'foreword', structure['foreword'])
            book.add_item(foreword)
            chapters.append(foreword)
            toc_entries.append(epub.Link(foreword.file_name, 'Foreword', 'foreword'))

        # Introduction
        if structure['introduction']:
            intro = self._create_section('Introduction', 'introduction', structure['introduction'])
            book.add_item(intro)
            chapters.append(intro)
            toc_entries.append(epub.Link(intro.file_name, 'Introduction', 'introduction'))

            # Introduction references
            if 'Introduction' in structure['references_by_chapter']:
                intro_refs = self._create_references_chapter(
                    'Introduction References',
                    'intro_refs',
                    structure['references_by_chapter']['Introduction']
                )
                book.add_item(intro_refs)
                chapters.append(intro_refs)

        # Chapters
        for ch_num in sorted(structure['chapters'].keys()):
            chapter_data = structure['chapters'][ch_num]

            # Chapter content
            chapter = self._create_chapter(ch_num, chapter_data)
            book.add_item(chapter)
            chapters.append(chapter)
            toc_entries.append(epub.Link(chapter.file_name, f'Chapter {ch_num}', f'chapter_{ch_num}'))

            # Chapter references
            if ch_num in structure['references_by_chapter']:
                ch_refs = self._create_references_chapter(
                    f'Chapter {ch_num} References',
                    f'chapter_{ch_num}_refs',
                    structure['references_by_chapter'][ch_num]
                )
                book.add_item(ch_refs)
                chapters.append(ch_refs)

        # Set TOC and spine
        book.toc = tuple(toc_entries)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ['nav'] + chapters

        # Write EPUB
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        epub.write_epub(str(output_file), book)

        print(f"✓ Generated: {output_file}")
        print(f"   Size: {output_file.stat().st_size:,} bytes")

        return output_file

    def _extract_structure(self) -> Dict:
        """Extract book structure from DOCX"""
        print("  📑 Extracting book structure...")

        structure = {
            'title': 'Caught in the Act',
            'foreword': [],
            'introduction': [],
            'chapters': {},
            'references_by_chapter': {}
        }

        current_section = None
        current_chapter = None
        in_references = False

        for para in self.doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Detect sections
            if re.match(r'^FOREWORD', text, re.IGNORECASE):
                current_section = 'foreword'
                current_chapter = None
                in_references = False
                continue

            elif re.match(r'^INTRODUCTION', text, re.IGNORECASE):
                current_section = 'introduction'
                current_chapter = None
                in_references = False
                continue

            elif re.match(r'^CHAPTER\s+(\d+)', text, re.IGNORECASE):
                match = re.match(r'^CHAPTER\s+(\d+)', text, re.IGNORECASE)
                chapter_num = int(match.group(1))
                current_section = 'chapters'
                current_chapter = chapter_num
                in_references = False

                if current_chapter not in structure['chapters']:
                    structure['chapters'][current_chapter] = {
                        'title': text,
                        'content': [],
                    }
                continue

            elif re.match(r'^(?:Introduction\s+)?References?\s*$', text, re.IGNORECASE):
                in_references = True
                continue

            # Store content
            if in_references:
                continue
            else:
                if current_section == 'foreword':
                    structure['foreword'].append(text)
                elif current_section == 'introduction':
                    structure['introduction'].append(text)
                elif current_section == 'chapters' and current_chapter:
                    structure['chapters'][current_chapter]['content'].append(text)

        # Organize references
        for ref_num, ref in self.references.items():
            chapter = ref['chapter']

            if chapter == 'Introduction':
                if 'Introduction' not in structure['references_by_chapter']:
                    structure['references_by_chapter']['Introduction'] = []
                structure['references_by_chapter']['Introduction'].append(ref)
            else:
                match = re.match(r'Chapter (\d+)', chapter)
                if match:
                    ch_num = int(match.group(1))
                    if ch_num not in structure['references_by_chapter']:
                        structure['references_by_chapter'][ch_num] = []
                    structure['references_by_chapter'][ch_num].append(ref)

        print(f"  ✓ Extracted complete book structure")

        return structure

    def _create_stylesheet(self):
        """Create CSS for EPUB"""
        css_content = '''
@namespace epub "http://www.idpf.org/2007/ops";

body {
    font-family: Georgia, serif;
    line-height: 1.8;
    margin: 1em;
    color: #333;
}

h1 {
    color: #2c3e50;
    font-size: 2em;
    margin-bottom: 0.5em;
    border-bottom: 3px solid #3498db;
    padding-bottom: 0.3em;
}

h2 {
    color: #2c3e50;
    font-size: 1.5em;
    margin-top: 1.5em;
    margin-bottom: 0.5em;
}

h3 {
    color: #2c3e50;
    font-size: 1.3em;
    margin-top: 1.2em;
    margin-bottom: 0.4em;
}

p {
    margin-bottom: 1em;
    text-align: justify;
}

.title-page {
    text-align: center;
    padding: 3em 1em;
}

.chapter-number {
    font-size: 1.1em;
    color: #666;
    text-transform: uppercase;
    letter-spacing: 2px;
    margin-bottom: 0.5em;
}

.reference {
    margin: 1.5em 0;
    padding: 1em;
    background: #f8f9fa;
    border-left: 4px solid #3498db;
    page-break-inside: avoid;
}

.ref-number {
    display: inline-block;
    background: #3498db;
    color: white;
    padding: 0.2em 0.6em;
    border-radius: 3px;
    font-weight: bold;
    margin-bottom: 0.5em;
}

.ref-bibliography {
    margin: 0.5em 0;
    line-height: 1.6;
}

.ref-relevance {
    margin: 1em 0;
    padding: 0.8em;
    background: white;
    border-left: 4px solid #e74c3c;
    font-size: 0.95em;
}

.relevance-label {
    font-weight: bold;
    color: #2c3e50;
    font-style: italic;
}

.ref-urls {
    margin-top: 1em;
    padding-top: 1em;
    border-top: 1px solid #ddd;
}

.url-link {
    display: block;
    margin: 0.5em 0;
    color: #3498db;
    text-decoration: none;
    word-wrap: break-word;
    font-size: 0.9em;
}
'''

        css = epub.EpubItem(
            uid="style",
            file_name="style.css",
            media_type="text/css",
            content=css_content
        )

        return css

    def _create_title_page(self):
        """Create title page"""
        content = f'''<?xml version='1.0' encoding='utf-8'?>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>Caught in the Act</title>
    <link href="style.css" rel="stylesheet" type="text/css"/>
</head>
<body>
    <div class="title-page">
        <h1>Caught in the Act</h1>
        <p style="font-size: 1.3em; margin: 1em 0;">Understanding Political Performance in the Digital Age</p>
        <p style="font-size: 1.2em; font-style: italic;">by Joe Ferguson</p>
        <p style="margin-top: 3em; color: #666;">Complete Edition with Enhanced References</p>
        <p style="color: #999;">{datetime.now().strftime('%B %d, %Y')}</p>
    </div>
</body>
</html>'''

        title = epub.EpubHtml(
            title='Title Page',
            file_name='title.xhtml',
            lang='en',
            content=content.encode('utf-8')
        )

        return title

    def _create_section(self, title: str, file_id: str, paragraphs: List[str]):
        """Create a section (Foreword/Introduction)"""
        paras_html = '\n'.join(f'<p>{self._esc(p)}</p>' for p in paragraphs)

        content = f'''<?xml version='1.0' encoding='utf-8'?>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{title}</title>
    <link href="style.css" rel="stylesheet" type="text/css"/>
</head>
<body>
    <h1>{title}</h1>
    {paras_html}
</body>
</html>'''

        section = epub.EpubHtml(
            title=title,
            file_name=f'{file_id}.xhtml',
            lang='en',
            content=content.encode('utf-8')
        )

        return section

    def _create_chapter(self, chapter_num: int, chapter_data: Dict):
        """Create a chapter"""
        paras_html = '\n'.join(f'<p>{self._esc(p)}</p>' for p in chapter_data['content'])

        content = f'''<?xml version='1.0' encoding='utf-8'?>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>Chapter {chapter_num}</title>
    <link href="style.css" rel="stylesheet" type="text/css"/>
</head>
<body>
    <div class="chapter-number">Chapter {chapter_num}</div>
    <h1>{chapter_data.get('title', f'Chapter {chapter_num}')}</h1>
    {paras_html}
</body>
</html>'''

        chapter = epub.EpubHtml(
            title=f'Chapter {chapter_num}',
            file_name=f'chapter_{chapter_num}.xhtml',
            lang='en',
            content=content.encode('utf-8')
        )

        return chapter

    def _create_references_chapter(self, title: str, file_id: str, references: List[Dict]):
        """Create references chapter"""
        refs_html = []

        for ref in sorted(references, key=lambda r: r['number']):
            ref_html = self._format_reference(ref)
            refs_html.append(ref_html)

        refs_content = '\n'.join(refs_html)

        content = f'''<?xml version='1.0' encoding='utf-8'?>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{title}</title>
    <link href="style.css" rel="stylesheet" type="text/css"/>
</head>
<body>
    <h2>{title}</h2>
    {refs_content}
</body>
</html>'''

        refs_chapter = epub.EpubHtml(
            title=title,
            file_name=f'{file_id}.xhtml',
            lang='en',
            content=content.encode('utf-8')
        )

        return refs_chapter

    def _format_reference(self, ref: Dict) -> str:
        """Format a reference"""
        bib = self._esc(ref['bibliography'])
        rel = self._esc(ref.get('relevance', ''))
        purl = ref.get('primary_url', '')
        surl = ref.get('secondary_url', '')

        rel_html = f'<div class="ref-relevance"><span class="relevance-label">Citation Narrative:</span> {rel}</div>' if rel else ''

        urls_html = ''
        if purl or surl:
            url_links = []
            if purl:
                url_links.append(f'<a href="{purl}" class="url-link">🔗 Primary: {purl}</a>')
            if surl:
                url_links.append(f'<a href="{surl}" class="url-link">📎 Secondary: {surl}</a>')
            urls_html = f'<div class="ref-urls">{" ".join(url_links)}</div>'

        return f'''<div class="reference">
    <span class="ref-number">[{ref['number']}]</span>
    <div class="ref-bibliography">{bib}</div>
    {rel_html}
    {urls_html}
</div>'''

    def _esc(self, text: str) -> str:
        """Escape HTML"""
        if not text:
            return ''
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def main():
    """Main execution"""
    print("="*70)
    print("Full Book EPUB Converter - Caught in the Act")
    print("="*70)

    # Check manuscript
    manuscript_path = 'source/250827_Caught_In_The_Act_Kindle.docx'
    if not Path(manuscript_path).exists():
        print(f"\n❌ Manuscript not found: {manuscript_path}")
        return 1

    # Convert
    converter = FullBookEPUBConverter(
        docx_path=manuscript_path,
        references_json='data/references_master.json'
    )

    output_path = converter.convert('outputs/complete_book.epub')

    print(f"\n✅ Complete book EPUB generated!")
    print(f"   Output: {output_path}")
    print(f"   Open with e-reader to view")

    return 0


if __name__ == "__main__":
    exit(main())
