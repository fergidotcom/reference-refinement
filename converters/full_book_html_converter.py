#!/usr/bin/env python3
"""
Full Book HTML Converter for "Caught in the Act"
Integrates complete manuscript with enhanced reference system
"""

import json
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class FullBookHTMLConverter:
    """Convert complete manuscript to HTML with enhanced references"""

    def __init__(self, docx_path: str, references_json: str = 'data/references_master.json'):
        print(f"📚 Loading complete manuscript from {docx_path}...")
        self.doc = Document(docx_path)

        print(f"📖 Loading enhanced references from {references_json}...")
        with open(references_json, 'r', encoding='utf-8') as f:
            ref_data = json.load(f)

        self.references = {ref['number']: ref for ref in ref_data['references']}
        self.statistics = ref_data['statistics']

        print(f"✓ Loaded {len(self.doc.paragraphs)} paragraphs from manuscript")
        print(f"✓ Loaded {len(self.references)} enhanced references")

    def convert(self, output_path: str = 'outputs/complete_book.html'):
        """Convert complete book to HTML"""
        print(f"\n🎨 Converting complete book to HTML...")

        # Extract manuscript structure
        structure = self._extract_structure()

        # Build HTML
        html = self._build_html(structure)

        # Write output
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html)

        print(f"✓ Generated: {output_file}")
        print(f"   Size: {output_file.stat().st_size:,} bytes")

        return output_file

    def _extract_structure(self) -> Dict:
        """Extract book structure from DOCX"""
        print("  📑 Extracting book structure...")

        structure = {
            'title': 'Caught in the Act',
            'foreword': [],
            'introduction': [],
            'chapters': {},
            'references_by_chapter': {}
        }

        current_section = None
        current_chapter = None
        in_references = False

        for para in self.doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Detect sections
            if re.match(r'^FOREWORD', text, re.IGNORECASE):
                current_section = 'foreword'
                current_chapter = None
                in_references = False
                continue

            elif re.match(r'^INTRODUCTION', text, re.IGNORECASE):
                current_section = 'introduction'
                current_chapter = None
                in_references = False
                continue

            elif re.match(r'^CHAPTER\s+(\d+)', text, re.IGNORECASE):
                match = re.match(r'^CHAPTER\s+(\d+)', text, re.IGNORECASE)
                chapter_num = int(match.group(1))
                current_section = 'chapters'
                current_chapter = chapter_num
                in_references = False

                if current_chapter not in structure['chapters']:
                    structure['chapters'][current_chapter] = {
                        'title': text,
                        'content': [],
                        'references': []
                    }
                continue

            # Detect reference sections
            elif re.match(r'^(?:Introduction\s+)?References?\s*$', text, re.IGNORECASE):
                in_references = True
                if current_section == 'introduction':
                    structure['references_by_chapter']['Introduction'] = []
                elif current_chapter:
                    if current_chapter not in structure['references_by_chapter']:
                        structure['references_by_chapter'][current_chapter] = []
                continue

            # Store content
            if in_references:
                # Skip old references - we'll use enhanced ones
                continue
            else:
                # Add to appropriate section
                if current_section == 'foreword':
                    structure['foreword'].append(text)
                elif current_section == 'introduction':
                    structure['introduction'].append(text)
                elif current_section == 'chapters' and current_chapter:
                    structure['chapters'][current_chapter]['content'].append(text)

        # Organize enhanced references by chapter
        for ref_num, ref in self.references.items():
            chapter = ref['chapter']

            if chapter == 'Introduction':
                if 'Introduction' not in structure['references_by_chapter']:
                    structure['references_by_chapter']['Introduction'] = []
                structure['references_by_chapter']['Introduction'].append(ref)
            else:
                # Extract chapter number from "Chapter N"
                match = re.match(r'Chapter (\d+)', chapter)
                if match:
                    ch_num = int(match.group(1))
                    if ch_num not in structure['references_by_chapter']:
                        structure['references_by_chapter'][ch_num] = []
                    structure['references_by_chapter'][ch_num].append(ref)

        print(f"  ✓ Found {len(structure['foreword'])} foreword paragraphs")
        print(f"  ✓ Found {len(structure['introduction'])} introduction paragraphs")
        print(f"  ✓ Found {len(structure['chapters'])} chapters")
        print(f"  ✓ Organized {len(self.references)} references")

        return structure

    def _build_html(self, structure: Dict) -> str:
        """Build complete HTML document"""
        css = self._generate_css()
        content = self._generate_content(structure)

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Caught in the Act - Complete Book</title>
    <style>
{css}
    </style>
</head>
<body>
    <div class="book-container">
{content}
    </div>

    <script>
        // Add smooth scrolling and navigation
        document.querySelectorAll('a[href^="#"]').forEach(anchor => {{
            anchor.addEventListener('click', function (e) {{
                e.preventDefault();
                const target = document.querySelector(this.getAttribute('href'));
                if (target) {{
                    target.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
                }}
            }});
        }});

        // Reference link tracking
        console.log('Caught in the Act - Complete Book loaded');
        console.log('Total references: ' + document.querySelectorAll('.reference-item').length);
    </script>
</body>
</html>"""

    def _generate_css(self) -> str:
        """Generate CSS for complete book"""
        return """
        @page {
            margin: 1in;
        }

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: Georgia, "Times New Roman", serif;
            line-height: 1.8;
            color: #1a1a1a;
            background: #f5f5f5;
            font-size: 16px;
        }

        .book-container {
            max-width: 800px;
            margin: 0 auto;
            background: white;
            padding: 60px 80px;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }

        /* Title Page */
        .title-page {
            text-align: center;
            padding: 100px 0;
            page-break-after: always;
        }

        .book-title {
            font-size: 3em;
            font-weight: bold;
            margin-bottom: 0.5em;
            color: #2c3e50;
        }

        .book-subtitle {
            font-size: 1.5em;
            color: #666;
            margin-bottom: 2em;
        }

        .book-author {
            font-size: 1.3em;
            font-style: italic;
            color: #333;
        }

        /* Table of Contents */
        .toc {
            page-break-after: always;
            margin: 60px 0;
        }

        .toc h2 {
            font-size: 2em;
            margin-bottom: 1em;
            border-bottom: 2px solid #2c3e50;
            padding-bottom: 0.3em;
        }

        .toc ul {
            list-style: none;
            padding: 0;
        }

        .toc li {
            margin: 0.8em 0;
            padding-left: 1em;
        }

        .toc a {
            color: #2c3e50;
            text-decoration: none;
            border-bottom: 1px dotted #999;
        }

        .toc a:hover {
            color: #3498db;
            border-bottom-color: #3498db;
        }

        /* Sections */
        .section {
            page-break-before: always;
            margin: 60px 0;
        }

        .section-title {
            font-size: 2.5em;
            margin-bottom: 1em;
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 0.3em;
        }

        .chapter {
            page-break-before: always;
            margin: 60px 0;
        }

        .chapter-number {
            font-size: 1.2em;
            color: #666;
            text-transform: uppercase;
            letter-spacing: 2px;
        }

        .chapter-title {
            font-size: 2.5em;
            margin: 0.3em 0 1em 0;
            color: #2c3e50;
        }

        /* Content */
        .content p {
            margin-bottom: 1.2em;
            text-align: justify;
            text-indent: 2em;
        }

        .content p:first-child {
            text-indent: 0;
        }

        /* References Section */
        .references-section {
            margin-top: 60px;
            padding-top: 30px;
            border-top: 2px solid #ccc;
        }

        .references-title {
            font-size: 2em;
            margin-bottom: 1em;
            color: #2c3e50;
        }

        .reference-item {
            margin: 1.5em 0;
            padding: 1em;
            background: #f9f9f9;
            border-left: 4px solid #3498db;
            page-break-inside: avoid;
        }

        .ref-number {
            display: inline-block;
            background: #3498db;
            color: white;
            padding: 0.2em 0.6em;
            border-radius: 3px;
            font-weight: bold;
            margin-right: 0.5em;
            font-size: 0.9em;
        }

        .ref-bibliography {
            margin: 0.5em 0;
            line-height: 1.6;
        }

        .ref-relevance {
            margin: 0.8em 0;
            padding: 0.8em;
            background: white;
            border-left: 3px solid #e74c3c;
            font-size: 0.95em;
            line-height: 1.6;
        }

        .relevance-label {
            font-weight: bold;
            color: #2c3e50;
            font-style: italic;
        }

        .ref-urls {
            margin-top: 0.8em;
            padding-top: 0.8em;
            border-top: 1px solid #ddd;
        }

        .ref-url {
            display: block;
            margin: 0.4em 0;
            color: #3498db;
            text-decoration: none;
            font-size: 0.9em;
            word-wrap: break-word;
        }

        .ref-url:hover {
            color: #2980b9;
            text-decoration: underline;
        }

        .ref-url::before {
            content: "🔗 ";
        }

        .ref-url.secondary::before {
            content: "📎 ";
        }

        /* Print styles */
        @media print {
            .book-container {
                max-width: 100%;
                padding: 0;
                box-shadow: none;
            }

            .section, .chapter {
                page-break-before: always;
            }

            .reference-item {
                page-break-inside: avoid;
            }

            a {
                color: #000;
            }

            .ref-url::after {
                content: " (" attr(href) ")";
                font-size: 0.8em;
                color: #666;
            }
        }

        @media screen and (max-width: 768px) {
            .book-container {
                padding: 30px 20px;
            }

            .book-title {
                font-size: 2em;
            }

            .section-title, .chapter-title {
                font-size: 1.8em;
            }
        }
"""

    def _generate_content(self, structure: Dict) -> str:
        """Generate book content HTML"""
        sections = []

        # Title page
        title_page = f"""        <div class="title-page">
            <h1 class="book-title">{structure['title']}</h1>
            <p class="book-subtitle">Understanding Political Performance in the Digital Age</p>
            <p class="book-author">by Joe Ferguson</p>
            <p style="margin-top: 3em; color: #666;">Complete Edition with Enhanced References</p>
            <p style="color: #999;">Generated {datetime.now().strftime('%B %d, %Y')}</p>
        </div>"""
        sections.append(title_page)

        # Table of contents
        toc = self._generate_toc(structure)
        sections.append(toc)

        # Foreword
        if structure['foreword']:
            foreword_html = self._generate_section('Foreword', structure['foreword'])
            sections.append(foreword_html)

        # Introduction
        if structure['introduction']:
            intro_html = self._generate_section('Introduction', structure['introduction'])
            sections.append(intro_html)

            # Introduction references
            if 'Introduction' in structure['references_by_chapter']:
                intro_refs = self._generate_references('Introduction', structure['references_by_chapter']['Introduction'])
                sections.append(intro_refs)

        # Chapters
        for ch_num in sorted(structure['chapters'].keys()):
            chapter = structure['chapters'][ch_num]
            chapter_html = self._generate_chapter(ch_num, chapter)
            sections.append(chapter_html)

            # Chapter references
            if ch_num in structure['references_by_chapter']:
                ch_refs = self._generate_references(f'Chapter {ch_num}', structure['references_by_chapter'][ch_num])
                sections.append(ch_refs)

        return '\n'.join(sections)

    def _generate_toc(self, structure: Dict) -> str:
        """Generate table of contents"""
        items = []

        items.append('<li><a href="#foreword">Foreword</a></li>')
        items.append('<li><a href="#introduction">Introduction</a></li>')

        for ch_num in sorted(structure['chapters'].keys()):
            items.append(f'<li><a href="#chapter-{ch_num}">Chapter {ch_num}</a></li>')

        toc_list = '\n            '.join(items)

        return f"""        <div class="toc">
            <h2>Table of Contents</h2>
            <ul>
            {toc_list}
            </ul>
        </div>"""

    def _generate_section(self, title: str, paragraphs: List[str]) -> str:
        """Generate a section (Foreword/Introduction)"""
        section_id = title.lower()

        content_paras = '\n            '.join(
            f'<p>{self._escape_html(p)}</p>' for p in paragraphs
        )

        return f"""        <div class="section" id="{section_id}">
            <h2 class="section-title">{title}</h2>
            <div class="content">
            {content_paras}
            </div>
        </div>"""

    def _generate_chapter(self, chapter_num: int, chapter: Dict) -> str:
        """Generate a chapter"""
        content_paras = '\n            '.join(
            f'<p>{self._escape_html(p)}</p>' for p in chapter['content']
        )

        return f"""        <div class="chapter" id="chapter-{chapter_num}">
            <div class="chapter-number">Chapter {chapter_num}</div>
            <h2 class="chapter-title">{chapter.get('title', f'Chapter {chapter_num}')}</h2>
            <div class="content">
            {content_paras}
            </div>
        </div>"""

    def _generate_references(self, section_name: str, references: List[Dict]) -> str:
        """Generate references section"""
        ref_items = []

        for ref in sorted(references, key=lambda r: r['number']):
            ref_html = self._format_reference(ref)
            ref_items.append(ref_html)

        refs_content = '\n'.join(ref_items)
        section_id = section_name.lower().replace(' ', '-')

        return f"""        <div class="references-section" id="refs-{section_id}">
            <h3 class="references-title">{section_name} References</h3>
{refs_content}
        </div>"""

    def _format_reference(self, ref: Dict) -> str:
        """Format a single reference"""
        ref_num = ref['number']
        bibliography = self._escape_html(ref['bibliography'])
        relevance = ref.get('relevance', '')
        primary_url = ref.get('primary_url')
        secondary_url = ref.get('secondary_url')

        # Relevance (Citation Narrative)
        relevance_html = ''
        if relevance:
            relevance_html = f"""
            <div class="ref-relevance">
                <span class="relevance-label">Citation Narrative:</span> {self._escape_html(relevance)}
            </div>"""

        # URLs
        urls_html = ''
        if primary_url or secondary_url:
            url_links = []

            if primary_url:
                url_links.append(f'<a href="{primary_url}" class="ref-url" target="_blank" rel="noopener">Primary Source: {primary_url}</a>')

            if secondary_url:
                url_links.append(f'<a href="{secondary_url}" class="ref-url secondary" target="_blank" rel="noopener">Secondary Source: {secondary_url}</a>')

            urls_html = f"""
            <div class="ref-urls">
                {chr(10).join('                ' + link for link in url_links)}
            </div>"""

        return f"""            <div class="reference-item">
                <span class="ref-number">[{ref_num}]</span>
                <div class="ref-bibliography">{bibliography}</div>
{relevance_html}
{urls_html}
            </div>"""

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        if not text:
            return ''

        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))


def main():
    """Main execution"""
    print("="*70)
    print("Full Book HTML Converter - Caught in the Act")
    print("="*70)

    # Check if manuscript exists
    manuscript_path = 'source/250827_Caught_In_The_Act_Kindle.docx'
    if not Path(manuscript_path).exists():
        print(f"\n❌ Manuscript not found: {manuscript_path}")
        print("   Please ensure the complete manuscript DOCX is in the source/ directory")
        return 1

    # Convert
    converter = FullBookHTMLConverter(
        docx_path=manuscript_path,
        references_json='data/references_master.json'
    )

    output_path = converter.convert('outputs/complete_book.html')

    print(f"\n✅ Complete book HTML generated!")
    print(f"   Output: {output_path}")
    print(f"   Open in browser to view the complete book")

    return 0


if __name__ == "__main__":
    exit(main())
